"""
Report Writer MCP Server (HTTP+SSE Transport)
Generates investigation reports in multiple formats (Markdown, PDF, DOCX).

Like Claude Code's ability to write .md files, this MCP server allows Claude-Flow to:
- Generate structured investigation reports
- Export to court-ready formats (PDF, DOCX)
- Create real-time markdown documentation
- Build timeline visualizations
- Produce Brady material summaries

Implements proper MCP SSE transport (protocol 2024-11-05):
- GET /sse - SSE endpoint, sends 'endpoint' event with POST URL
- POST /messages - Client sends JSON-RPC requests here
- Server responds via SSE with matching request IDs

Tools:
- create_report: Start new investigation report
- add_section: Add section to report
- add_finding: Add investigation finding
- add_contradiction: Add contradiction entry
- add_brady_item: Add Brady material to report
- add_timeline_event: Add event to timeline
- add_citation: Add verified citation
- export_markdown: Export as .md file
- export_pdf: Export as PDF
- export_docx: Export as Word document
- get_report_status: Get report contents
- list_reports: List all reports for case
"""

import os
import json
import asyncio
import uuid
from typing import Optional, Dict, Any, List
from datetime import datetime
from pathlib import Path
from aiohttp import web

# Configuration
SERVER_PORT = int(os.environ.get("MCP_SERVER_PORT", "3015"))
AUTH_TOKEN = os.environ.get("MCP_AUTH_TOKEN", "")
REPORTS_DIR = Path(os.environ.get("REPORTS_OUTPUT_DIR", "/data/reports"))
TEMPLATES_DIR = Path(os.environ.get("TEMPLATES_DIR", "/templates"))

# Ensure directories exist
REPORTS_DIR.mkdir(parents=True, exist_ok=True)

# Store active SSE sessions
SSE_SESSIONS: Dict[str, Any] = {}

# In-memory report storage (would be persisted to PostgreSQL in production)
REPORTS: Dict[str, Dict] = {}


# Tool definitions for MCP
TOOLS = [
    {
        "name": "create_report",
        "description": "Create a new investigation report. Returns report ID for adding content.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "case_id": {"type": "string", "description": "Case ID (e.g., 'us-v-blackman')"},
                "report_type": {
                    "type": "string",
                    "description": "Type of report",
                    "enum": ["investigation_summary", "witness_analysis", "brady_report",
                             "timeline_report", "contradiction_report", "financial_analysis",
                             "cross_examination_prep"]
                },
                "title": {"type": "string", "description": "Report title"}
            },
            "required": ["case_id", "report_type", "title"]
        }
    },
    {
        "name": "add_section",
        "description": "Add a section to an existing report.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "report_id": {"type": "string", "description": "Report ID"},
                "heading": {"type": "string", "description": "Section heading"},
                "content": {"type": "string", "description": "Section content (Markdown supported)"},
                "level": {"type": "integer", "description": "Heading level 1-4 (default 2)", "default": 2}
            },
            "required": ["report_id", "heading", "content"]
        }
    },
    {
        "name": "add_finding",
        "description": "Add an investigation finding to a report.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "report_id": {"type": "string", "description": "Report ID"},
                "finding": {"type": "string", "description": "Finding text"},
                "sources": {"type": "array", "items": {"type": "string"}, "description": "Source document IDs"},
                "confidence": {"type": "string", "description": "Confidence level: HIGH, MEDIUM, LOW", "default": "MEDIUM"}
            },
            "required": ["report_id", "finding"]
        }
    },
    {
        "name": "add_contradiction",
        "description": "Document a contradiction found between sources.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "report_id": {"type": "string", "description": "Report ID"},
                "source1": {"type": "object", "description": "First source {doc: string, quote: string}"},
                "source2": {"type": "object", "description": "Second source {doc: string, quote: string}"},
                "description": {"type": "string", "description": "Description of the contradiction"},
                "severity": {"type": "string", "description": "Severity: HIGH, MEDIUM, LOW", "default": "MEDIUM"}
            },
            "required": ["report_id", "source1", "source2", "description"]
        }
    },
    {
        "name": "add_brady_item",
        "description": "Add Brady/Giglio material to a report.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "report_id": {"type": "string", "description": "Report ID"},
                "item_type": {"type": "string", "description": "Type: exculpatory, impeachment, deal, coercion"},
                "description": {"type": "string", "description": "Description of Brady material"},
                "sources": {"type": "array", "items": {"type": "string"}, "description": "Source document IDs"},
                "priority": {"type": "string", "description": "Priority: HIGH, MEDIUM, LOW", "default": "MEDIUM"}
            },
            "required": ["report_id", "item_type", "description"]
        }
    },
    {
        "name": "add_timeline_event",
        "description": "Add a timeline event to a report.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "report_id": {"type": "string", "description": "Report ID"},
                "date": {"type": "string", "description": "Date (YYYY-MM-DD or YYYY-MM-DD HH:MM)"},
                "event": {"type": "string", "description": "Event description"},
                "source": {"type": "string", "description": "Source document ID"},
                "actors": {"type": "array", "items": {"type": "string"}, "description": "People involved"}
            },
            "required": ["report_id", "date", "event"]
        }
    },
    {
        "name": "add_citation",
        "description": "Add a verified legal citation to a report.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "report_id": {"type": "string", "description": "Report ID"},
                "citation": {"type": "string", "description": "Legal citation (e.g., 'Brady v. Maryland, 373 U.S. 83')"},
                "context": {"type": "string", "description": "How this citation is relevant"},
                "verified": {"type": "boolean", "description": "Whether citation was verified", "default": True}
            },
            "required": ["report_id", "citation", "context"]
        }
    },
    {
        "name": "export_markdown",
        "description": "Export report as Markdown file.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "report_id": {"type": "string", "description": "Report ID"},
                "output_path": {"type": "string", "description": "Optional output path (defaults to reports dir)"}
            },
            "required": ["report_id"]
        }
    },
    {
        "name": "export_pdf",
        "description": "Export report as PDF file.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "report_id": {"type": "string", "description": "Report ID"},
                "output_path": {"type": "string", "description": "Optional output path"}
            },
            "required": ["report_id"]
        }
    },
    {
        "name": "export_docx",
        "description": "Export report as Word document.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "report_id": {"type": "string", "description": "Report ID"},
                "output_path": {"type": "string", "description": "Optional output path"}
            },
            "required": ["report_id"]
        }
    },
    {
        "name": "get_report_status",
        "description": "Get current report contents and structure.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "report_id": {"type": "string", "description": "Report ID"}
            },
            "required": ["report_id"]
        }
    },
    {
        "name": "list_reports",
        "description": "List all reports for a case.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "case_id": {"type": "string", "description": "Case ID"}
            },
            "required": ["case_id"]
        }
    }
]


# ============================================================================
# Tool Implementations
# ============================================================================

async def create_report(case_id: str, report_type: str, title: str) -> str:
    """Create a new report."""
    report_id = str(uuid.uuid4())[:8]

    REPORTS[report_id] = {
        "id": report_id,
        "case_id": case_id,
        "type": report_type,
        "title": title,
        "created_at": datetime.utcnow().isoformat(),
        "updated_at": datetime.utcnow().isoformat(),
        "sections": [],
        "findings": [],
        "contradictions": [],
        "brady_items": [],
        "timeline": [],
        "citations": []
    }

    return f"**Report Created**\n\n**ID:** `{report_id}`\n**Type:** {report_type}\n**Title:** {title}"


async def add_section(report_id: str, heading: str, content: str, level: int = 2) -> str:
    """Add a section to report."""
    if report_id not in REPORTS:
        return f"Report not found: {report_id}"

    section_id = str(uuid.uuid4())[:8]
    REPORTS[report_id]["sections"].append({
        "id": section_id,
        "heading": heading,
        "content": content,
        "level": level
    })
    REPORTS[report_id]["updated_at"] = datetime.utcnow().isoformat()

    return f"**Section Added**\n\n**ID:** `{section_id}`\n**Heading:** {heading}"


async def add_finding(report_id: str, finding: str, sources: List[str] = None,
                       confidence: str = "MEDIUM") -> str:
    """Add a finding to report."""
    if report_id not in REPORTS:
        return f"Report not found: {report_id}"

    finding_id = str(uuid.uuid4())[:8]
    REPORTS[report_id]["findings"].append({
        "id": finding_id,
        "finding": finding,
        "sources": sources or [],
        "confidence": confidence
    })
    REPORTS[report_id]["updated_at"] = datetime.utcnow().isoformat()

    return f"**Finding Added**\n\n**ID:** `{finding_id}`\n**Confidence:** {confidence}"


async def add_contradiction(report_id: str, source1: dict, source2: dict,
                            description: str, severity: str = "MEDIUM") -> str:
    """Add a contradiction to report."""
    if report_id not in REPORTS:
        return f"Report not found: {report_id}"

    contradiction_id = str(uuid.uuid4())[:8]
    REPORTS[report_id]["contradictions"].append({
        "id": contradiction_id,
        "source1": source1,
        "source2": source2,
        "description": description,
        "severity": severity
    })
    REPORTS[report_id]["updated_at"] = datetime.utcnow().isoformat()

    return f"**Contradiction Documented**\n\n**ID:** `{contradiction_id}`\n**Severity:** {severity}"


async def add_brady_item(report_id: str, item_type: str, description: str,
                          sources: List[str] = None, priority: str = "MEDIUM") -> str:
    """Add Brady item to report."""
    if report_id not in REPORTS:
        return f"Report not found: {report_id}"

    item_id = str(uuid.uuid4())[:8]
    REPORTS[report_id]["brady_items"].append({
        "id": item_id,
        "type": item_type,
        "description": description,
        "sources": sources or [],
        "priority": priority
    })
    REPORTS[report_id]["updated_at"] = datetime.utcnow().isoformat()

    return f"**Brady Item Added**\n\n**ID:** `{item_id}`\n**Type:** {item_type}\n**Priority:** {priority}"


async def add_timeline_event(report_id: str, date: str, event: str,
                              source: str = None, actors: List[str] = None) -> str:
    """Add timeline event to report."""
    if report_id not in REPORTS:
        return f"Report not found: {report_id}"

    event_id = str(uuid.uuid4())[:8]
    REPORTS[report_id]["timeline"].append({
        "id": event_id,
        "date": date,
        "event": event,
        "source": source,
        "actors": actors or []
    })
    REPORTS[report_id]["updated_at"] = datetime.utcnow().isoformat()

    # Sort timeline by date
    REPORTS[report_id]["timeline"].sort(key=lambda x: x["date"])

    return f"**Timeline Event Added**\n\n**ID:** `{event_id}`\n**Date:** {date}"


async def add_citation(report_id: str, citation: str, context: str,
                        verified: bool = True) -> str:
    """Add citation to report."""
    if report_id not in REPORTS:
        return f"Report not found: {report_id}"

    citation_id = str(uuid.uuid4())[:8]
    REPORTS[report_id]["citations"].append({
        "id": citation_id,
        "citation": citation,
        "context": context,
        "verified": verified
    })
    REPORTS[report_id]["updated_at"] = datetime.utcnow().isoformat()

    verified_str = "✅ Verified" if verified else "⚠️ Unverified"
    return f"**Citation Added** {verified_str}\n\n**ID:** `{citation_id}`\n**Citation:** {citation}"


def _generate_markdown(report: dict) -> str:
    """Generate markdown content for a report."""
    md = []

    # Title and metadata
    md.append(f"# {report['title']}\n")
    md.append(f"**Case:** {report['case_id']}")
    md.append(f"**Type:** {report['type']}")
    md.append(f"**Generated:** {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}")
    md.append("")
    md.append("---")
    md.append("")

    # Sections
    for section in report.get("sections", []):
        level = section.get("level", 2)
        md.append(f"{'#' * level} {section['heading']}\n")
        md.append(section['content'])
        md.append("")

    # Findings
    if report.get("findings"):
        md.append("## Key Findings\n")
        for i, finding in enumerate(report["findings"], 1):
            confidence_icon = {"HIGH": "🔴", "MEDIUM": "🟡", "LOW": "🟢"}.get(finding["confidence"], "⚪")
            md.append(f"{i}. {confidence_icon} **{finding['confidence']}** - {finding['finding']}")
            if finding.get("sources"):
                md.append(f"   - Sources: {', '.join(finding['sources'])}")
        md.append("")

    # Contradictions
    if report.get("contradictions"):
        md.append("## Contradictions Found\n")
        for i, contra in enumerate(report["contradictions"], 1):
            severity_icon = {"HIGH": "🔴", "MEDIUM": "🟡", "LOW": "🟢"}.get(contra["severity"], "⚪")
            md.append(f"### {i}. {severity_icon} {contra['description']}\n")
            md.append(f"**Source 1 ({contra['source1'].get('doc', 'Unknown')}):**")
            md.append(f"> {contra['source1'].get('quote', 'N/A')}\n")
            md.append(f"**Source 2 ({contra['source2'].get('doc', 'Unknown')}):**")
            md.append(f"> {contra['source2'].get('quote', 'N/A')}\n")
        md.append("")

    # Brady Items
    if report.get("brady_items"):
        md.append("## Brady/Giglio Material\n")
        md.append("| Priority | Type | Description | Sources |")
        md.append("|----------|------|-------------|---------|")
        for item in report["brady_items"]:
            priority_icon = {"HIGH": "🔴", "MEDIUM": "🟡", "LOW": "🟢"}.get(item["priority"], "⚪")
            sources = ", ".join(item.get("sources", [])) or "N/A"
            md.append(f"| {priority_icon} {item['priority']} | {item['type']} | {item['description'][:50]}... | {sources} |")
        md.append("")

    # Timeline
    if report.get("timeline"):
        md.append("## Timeline\n")
        for event in report["timeline"]:
            actors_str = f" ({', '.join(event['actors'])})" if event.get("actors") else ""
            source_str = f" [Source: {event['source']}]" if event.get("source") else ""
            md.append(f"- **{event['date']}:** {event['event']}{actors_str}{source_str}")
        md.append("")

    # Citations
    if report.get("citations"):
        md.append("## Legal Citations\n")
        for citation in report["citations"]:
            verified_str = "✅" if citation["verified"] else "⚠️"
            md.append(f"- {verified_str} **{citation['citation']}** - {citation['context']}")
        md.append("")

    # Footer
    md.append("---")
    md.append("*Generated by LegalAI Report Writer MCP*")

    return "\n".join(md)


async def export_markdown(report_id: str, output_path: str = None) -> str:
    """Export report as Markdown."""
    if report_id not in REPORTS:
        return f"Report not found: {report_id}"

    report = REPORTS[report_id]
    md_content = _generate_markdown(report)

    # Determine output path
    if output_path:
        file_path = Path(output_path)
    else:
        safe_title = "".join(c if c.isalnum() or c in "-_ " else "" for c in report["title"])
        safe_title = safe_title.replace(" ", "-").lower()[:50]
        file_path = REPORTS_DIR / report["case_id"] / f"{safe_title}.md"

    file_path.parent.mkdir(parents=True, exist_ok=True)
    file_path.write_text(md_content, encoding="utf-8")

    return f"**Markdown Exported**\n\n**Path:** `{file_path}`\n**Size:** {len(md_content):,} bytes"


async def export_pdf(report_id: str, output_path: str = None) -> str:
    """Export report as PDF."""
    if report_id not in REPORTS:
        return f"Report not found: {report_id}"

    report = REPORTS[report_id]

    # First generate markdown
    md_content = _generate_markdown(report)

    # Determine output path
    if output_path:
        file_path = Path(output_path)
    else:
        safe_title = "".join(c if c.isalnum() or c in "-_ " else "" for c in report["title"])
        safe_title = safe_title.replace(" ", "-").lower()[:50]
        file_path = REPORTS_DIR / report["case_id"] / f"{safe_title}.pdf"

    file_path.parent.mkdir(parents=True, exist_ok=True)

    # Try to convert to PDF using available tools
    try:
        # Try pandoc first
        import subprocess
        md_path = file_path.with_suffix(".md")
        md_path.write_text(md_content, encoding="utf-8")

        result = subprocess.run(
            ["pandoc", str(md_path), "-o", str(file_path), "--pdf-engine=xelatex"],
            capture_output=True,
            text=True
        )

        if result.returncode == 0:
            md_path.unlink()  # Clean up temp markdown
            return f"**PDF Exported**\n\n**Path:** `{file_path}`"
        else:
            # Pandoc failed, save as markdown instead
            return f"**PDF conversion failed** (pandoc not available)\n\nMarkdown saved to: `{md_path}`\n\nError: {result.stderr[:200]}"

    except FileNotFoundError:
        # Pandoc not installed, save markdown as fallback
        md_path = file_path.with_suffix(".md")
        md_path.write_text(md_content, encoding="utf-8")
        return f"**PDF export unavailable** (pandoc not installed)\n\nMarkdown saved to: `{md_path}`"
    except Exception as e:
        return f"Error exporting PDF: {str(e)}"


async def export_docx(report_id: str, output_path: str = None) -> str:
    """Export report as Word document."""
    if report_id not in REPORTS:
        return f"Report not found: {report_id}"

    report = REPORTS[report_id]

    # Determine output path
    if output_path:
        file_path = Path(output_path)
    else:
        safe_title = "".join(c if c.isalnum() or c in "-_ " else "" for c in report["title"])
        safe_title = safe_title.replace(" ", "-").lower()[:50]
        file_path = REPORTS_DIR / report["case_id"] / f"{safe_title}.docx"

    file_path.parent.mkdir(parents=True, exist_ok=True)

    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # Title
        title = doc.add_heading(report["title"], 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Metadata
        doc.add_paragraph(f"Case: {report['case_id']}")
        doc.add_paragraph(f"Type: {report['type']}")
        doc.add_paragraph(f"Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}")
        doc.add_paragraph()

        # Sections
        for section in report.get("sections", []):
            doc.add_heading(section["heading"], level=section.get("level", 2))
            doc.add_paragraph(section["content"])

        # Findings
        if report.get("findings"):
            doc.add_heading("Key Findings", level=2)
            for i, finding in enumerate(report["findings"], 1):
                para = doc.add_paragraph()
                para.add_run(f"{i}. [{finding['confidence']}] ").bold = True
                para.add_run(finding["finding"])
                if finding.get("sources"):
                    doc.add_paragraph(f"   Sources: {', '.join(finding['sources'])}")

        # Brady Items
        if report.get("brady_items"):
            doc.add_heading("Brady/Giglio Material", level=2)
            for item in report["brady_items"]:
                para = doc.add_paragraph()
                para.add_run(f"[{item['priority']}] {item['type'].upper()}: ").bold = True
                para.add_run(item["description"])

        # Timeline
        if report.get("timeline"):
            doc.add_heading("Timeline", level=2)
            for event in report["timeline"]:
                para = doc.add_paragraph()
                para.add_run(f"{event['date']}: ").bold = True
                para.add_run(event["event"])

        doc.save(str(file_path))

        return f"**DOCX Exported**\n\n**Path:** `{file_path}`"

    except ImportError:
        # python-docx not installed, fall back to markdown
        md_content = _generate_markdown(report)
        md_path = file_path.with_suffix(".md")
        md_path.write_text(md_content, encoding="utf-8")
        return f"**DOCX export unavailable** (python-docx not installed)\n\nMarkdown saved to: `{md_path}`"
    except Exception as e:
        return f"Error exporting DOCX: {str(e)}"


async def get_report_status(report_id: str) -> str:
    """Get report status and contents."""
    if report_id not in REPORTS:
        return f"Report not found: {report_id}"

    report = REPORTS[report_id]

    output = f"**Report: {report['title']}**\n\n"
    output += f"**ID:** `{report_id}`\n"
    output += f"**Case:** {report['case_id']}\n"
    output += f"**Type:** {report['type']}\n"
    output += f"**Created:** {report['created_at']}\n"
    output += f"**Updated:** {report['updated_at']}\n\n"

    output += "**Contents:**\n"
    output += f"  - Sections: {len(report.get('sections', []))}\n"
    output += f"  - Findings: {len(report.get('findings', []))}\n"
    output += f"  - Contradictions: {len(report.get('contradictions', []))}\n"
    output += f"  - Brady Items: {len(report.get('brady_items', []))}\n"
    output += f"  - Timeline Events: {len(report.get('timeline', []))}\n"
    output += f"  - Citations: {len(report.get('citations', []))}\n"

    return output


async def list_reports(case_id: str) -> str:
    """List reports for a case."""
    case_reports = [r for r in REPORTS.values() if r["case_id"] == case_id]

    if not case_reports:
        return f"No reports found for case: {case_id}"

    output = f"**Reports for {case_id} ({len(case_reports)}):**\n\n"
    for report in case_reports:
        output += f"**{report['title']}**\n"
        output += f"   ID: `{report['id']}` | Type: {report['type']}\n"
        output += f"   Created: {report['created_at']}\n\n"

    return output


# Tool dispatcher
TOOL_HANDLERS = {
    "create_report": create_report,
    "add_section": add_section,
    "add_finding": add_finding,
    "add_contradiction": add_contradiction,
    "add_brady_item": add_brady_item,
    "add_timeline_event": add_timeline_event,
    "add_citation": add_citation,
    "export_markdown": export_markdown,
    "export_pdf": export_pdf,
    "export_docx": export_docx,
    "get_report_status": get_report_status,
    "list_reports": list_reports
}


# ============================================================================
# MCP Protocol Handlers
# ============================================================================

async def handle_mcp_request(data: dict, session_id: str = None) -> dict:
    """Handle a JSON-RPC MCP request and return response."""
    method = data.get("method", "")
    request_id = data.get("id")
    params = data.get("params", {})

    if method == "initialize":
        return {
            "jsonrpc": "2.0",
            "id": request_id,
            "result": {
                "protocolVersion": "2024-11-05",
                "capabilities": {"tools": {}},
                "serverInfo": {"name": "report-writer", "version": "1.0.0"}
            }
        }
    elif method == "tools/list":
        return {
            "jsonrpc": "2.0",
            "id": request_id,
            "result": {"tools": TOOLS}
        }
    elif method == "tools/call":
        tool_name = params.get("name")
        arguments = params.get("arguments", {})

        if tool_name not in TOOL_HANDLERS:
            return {
                "jsonrpc": "2.0",
                "id": request_id,
                "error": {"code": -32601, "message": f"Unknown tool: {tool_name}"}
            }

        handler = TOOL_HANDLERS[tool_name]
        result = await handler(**arguments)

        return {
            "jsonrpc": "2.0",
            "id": request_id,
            "result": {"content": [{"type": "text", "text": result}]}
        }
    elif method == "notifications/initialized":
        return None
    else:
        return {
            "jsonrpc": "2.0",
            "id": request_id,
            "error": {"code": -32601, "message": f"Unknown method: {method}"}
        }


# ============================================================================
# HTTP+SSE Endpoints
# ============================================================================

def verify_auth(request) -> bool:
    """Verify authorization token."""
    if not AUTH_TOKEN:
        return True

    auth_header = request.headers.get("Authorization", "")
    if auth_header.startswith("Bearer ") and auth_header[7:] == AUTH_TOKEN:
        return True
    if request.headers.get("api_key") == AUTH_TOKEN:
        return True
    if request.query.get("api_key") == AUTH_TOKEN:
        return True
    return False


async def handle_sse_get(request):
    """SSE endpoint for MCP protocol."""
    if not verify_auth(request):
        return web.json_response({"error": "Unauthorized"}, status=401)

    session_id = str(uuid.uuid4())

    response = web.StreamResponse(
        status=200,
        headers={
            'Content-Type': 'text/event-stream',
            'Cache-Control': 'no-cache',
            'Connection': 'keep-alive',
            'Access-Control-Allow-Origin': '*',
        }
    )
    await response.prepare(request)

    SSE_SESSIONS[session_id] = response

    try:
        await response.write(f"event: endpoint\ndata: /sse?session_id={session_id}\n\n".encode())

        while True:
            await asyncio.sleep(30)
            try:
                await response.write(f": ping\n\n".encode())
            except:
                break

    except asyncio.CancelledError:
        pass
    finally:
        SSE_SESSIONS.pop(session_id, None)

    return response


async def handle_sse_post(request):
    """Handle POST to /sse."""
    if not verify_auth(request):
        return web.json_response({"error": "Unauthorized"}, status=401)

    try:
        data = await request.json()
        session_id = request.query.get("session_id")

        response_data = await handle_mcp_request(data, session_id)

        if response_data and session_id and session_id in SSE_SESSIONS:
            sse_response = SSE_SESSIONS.get(session_id)
            if sse_response:
                await sse_response.write(f"event: message\ndata: {json.dumps(response_data)}\n\n".encode())
            return web.Response(status=202)
        elif response_data:
            return web.json_response(response_data)
        else:
            return web.Response(status=202)

    except Exception as e:
        return web.json_response(
            {"jsonrpc": "2.0", "id": None, "error": {"code": -32603, "message": str(e)}},
            status=500
        )


async def handle_sse(request):
    """Route /sse based on HTTP method."""
    if request.method == 'GET':
        return await handle_sse_get(request)
    elif request.method == 'POST':
        return await handle_sse_post(request)
    return web.Response(status=405)


async def handle_health(request):
    """Health check endpoint."""
    return web.json_response({
        "status": "healthy",
        "service": "report-writer-mcp",
        "tools": len(TOOLS),
        "reports_dir": str(REPORTS_DIR),
        "active_reports": len(REPORTS),
        "active_sessions": len(SSE_SESSIONS)
    })


async def handle_cors_preflight(request):
    """Handle CORS preflight."""
    return web.Response(
        status=200,
        headers={
            'Access-Control-Allow-Origin': '*',
            'Access-Control-Allow-Methods': 'GET, POST, OPTIONS',
            'Access-Control-Allow-Headers': 'Content-Type, Authorization, api_key',
        }
    )


def create_app():
    """Create the aiohttp application."""
    app = web.Application()

    @web.middleware
    async def cors_middleware(request, handler):
        if request.method == 'OPTIONS':
            return await handle_cors_preflight(request)
        response = await handler(request)
        response.headers['Access-Control-Allow-Origin'] = '*'
        return response

    app.middlewares.append(cors_middleware)

    app.router.add_get("/sse", handle_sse)
    app.router.add_post("/sse", handle_sse)
    app.router.add_options("/sse", handle_cors_preflight)
    app.router.add_get("/health", handle_health)

    return app


if __name__ == "__main__":
    print("=" * 60)
    print("Report Writer MCP Server (SSE Transport)")
    print("=" * 60)
    print(f"Port: {SERVER_PORT}")
    print(f"Reports Directory: {REPORTS_DIR}")
    print(f"SSE endpoint: http://0.0.0.0:{SERVER_PORT}/sse")
    print(f"Health check: http://0.0.0.0:{SERVER_PORT}/health")
    print(f"Tools available: {len(TOOLS)}")
    for tool in TOOLS:
        print(f"  - {tool['name']}")
    print("=" * 60)

    app = create_app()
    web.run_app(app, host="0.0.0.0", port=SERVER_PORT)
